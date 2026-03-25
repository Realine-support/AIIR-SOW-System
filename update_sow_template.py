"""
update_sow_template_v2.py
Reads the original Client A.EC.Jonathan Bailey.docx and produces
AIIR_IGNITE_SOW_Template.docx with all {{PLACEHOLDER}} tokens correct.
"""
from docx import Document
import re

INPUT_PATH  = r"D:\AIIR\Client A.EC.Jonathan Bailey.docx"
OUTPUT_PATH = r"D:\AIIR\AIIR_IGNITE_SOW_Template.docx"

# ---------------------------------------------------------------------------
# Replacement map — order matters: more specific → less specific
# Each entry: (find_str, replace_str)
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    # ---- IDs & names -------------------------------------------------------
    ("395830485039589",          "{{HUBSPOT_ID}}"),
    ("Jonathan Kirschner",       "{{COACH_NAME}}"),

    # ---- Company (3 spec locations + header "Client:" line = 4 total OK) --
    ("Client A",                 "{{CLIENT_COMPANY}}"),

    # ---- Stakeholder count -------------------------------------------------
    ("four formal stakeholder meetings",
     "{{STAKEHOLDER_COUNT}} formal stakeholder meetings"),

    # ---- Developmental History: inject DEV_HISTORY_TEXT -------------------
    # Original: "Client and coach meet for a Developmental History Interview.
    #  This interview gathers essential baseline information..."
    ("Client and coach meet for a Developmental History Interview. This interview gathers essential baseline information about the client",
     "{{CLIENT_TERM}} and coach meet for a Developmental History Interview. {{DEV_HISTORY_TEXT}} This interview gathers essential baseline information about the {{CLIENT_TERM}}"),

    # ---- Interview count ---------------------------------------------------
    ("up to 8 confidential interviews with key raters",
     "{{INTERVIEW_COUNT}}"),

    # ---- Streams of data ---------------------------------------------------
    ("three streams of data",    "{{STREAMS_COUNT}} streams of data"),

    # ---- Assessment Feedback: full sentence replacement -------------------
    # Original full paragraph: "Client and coach meet for an in-depth discussion...
    #  This meeting is typically two-hours in duration."
    # Replace entire description sentence with placeholder
    ("Client and coach meet for an in-depth discussion of all the findings from the assessment process. This meeting is typically two-hours in duration.",
     "{{ASSESSMENT_FEEDBACK_TEXT}}"),

    # ---- Dev Planning: full sentence replacement --------------------------
    ("Client and coach schedule a 1-hour session a week after the Assessment Feedback session to solidify the Strategic Development Plan and prepare for the 2nd stakeholder meeting.",
     "{{DEV_PLANNING_TEXT}}"),

    # ---- Pricing -----------------------------------------------------------
    ("$27,445",                  "{{TOTAL_PRICE}}"),

    # ---- Payment structure (two bullet lines → single token each) ---------
    ("50% of the engagement fees will be invoiced upon kick off of the engagement",
     "{{PAYMENT_STRUCTURE}}"),
    ("50% of the engagement fees will be invoiced three months following the kick off of the engagement",
     ""),   # remove second line; Workflow 2 inserts full PAYMENT_STRUCTURE text

    # ---- Net days ----------------------------------------------------------
    ("within 45 days",           "within {{NET_DAYS}} days"),

    # ---- CLIENT_TERM replacements (all forms) ------------------------------
    # Body paragraph intro
    ("for the client to learn",  "for the {{CLIENT_TERM}} to learn"),
    # Table: Coaching Kickoff
    ("Coach and client meet for a brief 30-minute session",
     "Coach and {{CLIENT_TERM}} meet for a brief 30-minute session"),
    # Table: ASM
    ("Attended by the client, coach",
     "Attended by the {{CLIENT_TERM}}, coach"),
    # Table: 360 assessment description already handled by DEV_HISTORY above
    # Table: Psychometric
    ("The client takes four psychometric",
     "The {{CLIENT_TERM}} takes four psychometric"),
    # Table: Mid-point review - client's progress
    ("evaluate the client\u2019s progress",
     "evaluate the {{CLIENT_TERM}}\u2019s progress"),
    ("evaluate the client's progress",
     "evaluate the {{CLIENT_TERM}}'s progress"),
    # Table: Business Impact Survey
    ("the client is sent a final survey",
     "the {{CLIENT_TERM}} is sent a final survey"),
    # Table: CAP
    ("the client creates a Continuing Action Plan",
     "the {{CLIENT_TERM}} creates a Continuing Action Plan"),
    # Table: Stakeholder Meeting 2
    ("The client presents this Strategic Development Plan",
     "The {{CLIENT_TERM}} presents this Strategic Development Plan"),
    # Table: Wrap-Up
    ("Coach, client, and manager meet",
     "Coach, {{CLIENT_TERM}}, and manager meet"),
    # Table: Mid-point
    ("The coach, client, and stakeholders schedule",
     "The coach, {{CLIENT_TERM}}, and stakeholders schedule"),
    # Table: Implementation
    ("The coach and client utilize",
     "The coach and {{CLIENT_TERM}} utilize"),
    # Confidentiality policy
    ("shared with the client only",
     "shared with the {{CLIENT_TERM}} only"),
    ("all information shared by the client with the coach",
     "all information shared by the {{CLIENT_TERM}} with the coach"),
]


def fix_paragraph(para):
    """Apply all replacements to a paragraph (across all runs)."""
    full = "".join(r.text for r in para.runs)
    new  = full
    for find, replace in REPLACEMENTS:
        new = new.replace(find, replace)
    if new == full:
        return
    if para.runs:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""


def fix_tables(container):
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_paragraph(para)


def main():
    doc = Document(INPUT_PATH)

    for para in doc.paragraphs:
        fix_paragraph(para)
    fix_tables(doc)

    for section in doc.sections:
        for part in [section.header, section.footer]:
            if part:
                for para in part.paragraphs:
                    fix_paragraph(para)
                fix_tables(part)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
