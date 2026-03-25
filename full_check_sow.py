from docx import Document
import re
import collections

doc = Document(r"D:\AIIR\AIIR_IGNITE_SOW_Template.docx")

all_text_lines = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text_lines.append(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    all_text_lines.append(para.text)

for section in doc.sections:
    for part in [section.header, section.footer]:
        if part:
            for para in part.paragraphs:
                if para.text.strip():
                    all_text_lines.append("[HDR/FTR] " + para.text)

full_text = "\n".join(all_text_lines)

found = re.findall(r'\{\{[A-Z_0-9]+\}\}', full_text)
counts = collections.Counter(found)

expected = {
    "{{HUBSPOT_ID}}": 2,
    "{{STAKEHOLDER_COUNT}}": 1,
    "{{COACH_NAME}}": 1,
    "{{CLIENT_TERM}}": None,
    "{{DEV_HISTORY_TEXT}}": 1,
    "{{INTERVIEW_COUNT}}": 1,
    "{{STREAMS_COUNT}}": 1,
    "{{ASSESSMENT_FEEDBACK_TEXT}}": 1,
    "{{DEV_PLANNING_TEXT}}": 1,
    "{{TOTAL_PRICE}}": 1,
    "{{PAYMENT_STRUCTURE}}": None,
    "{{NET_DAYS}}": 1,
    "{{CLIENT_COMPANY}}": 3,
}

print("=" * 55)
print("PLACEHOLDER AUDIT vs SETUP_INSTRUCTIONS.md")
print("=" * 55)
for ph, exp_count in expected.items():
    n = counts.get(ph, 0)
    if exp_count is None:
        status = "OK    " if n > 0 else "MISS  "
        print(f"{status} {ph}  (found {n}x)")
    else:
        if n == exp_count:
            status = "OK    "
        elif n > 0:
            status = "WARN  "
        else:
            status = "MISS  "
        print(f"{status} {ph}  (expected {exp_count}x, found {n}x)")

print("\nAll tokens found in doc:")
for ph, n in sorted(counts.items()):
    flag = "" if ph in expected else "  <- EXTRA (not in spec)"
    print(f"  {ph} x{n}{flag}")

print("\n" + "=" * 55)
print("LEFTOVER CLIENT DATA CHECK")
print("=" * 55)
for b in ["Jonathan Bailey", "Jonathan Kirschner", "Client A", "$27,445", "395830485039589"]:
    flag = "STILL PRESENT" if b in full_text else "cleaned"
    print(f"{flag}  \"{b}\"")

print("\n" + "=" * 55)
print("FULL DOC TEXT")
print("=" * 55)
print(full_text)
