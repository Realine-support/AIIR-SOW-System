from docx import Document
import re, collections, sys

doc = Document(r"D:\AIIR\AIIR_IGNITE_SOW_Template.docx")

lines = []
for para in doc.paragraphs:
    lines.append(repr(para.text))
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                lines.append("[TABLE] " + repr(para.text))
for section in doc.sections:
    for part in [section.header, section.footer]:
        if part:
            for para in part.paragraphs:
                lines.append("[HDR/FTR] " + repr(para.text))

with open(r"D:\AIIR\sow_raw.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print("Written to sow_raw.txt")
